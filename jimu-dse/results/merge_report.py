# -*- coding: utf-8 -*-
"""Merge the planning markdown (平台化改造规划.md) into the existing optimization
report docx (NPU固件优化重构报告.docx), using the SAME visual style helpers as
build_report.py so the combined document reads as one system. Content of both
sources is preserved — the markdown is rendered in full (headings, tables, code
blocks, lists, blockquotes) and appended as a new major section.

Output: NPU固件优化与平台化改造报告.docx (a new file; originals untouched).
"""
import re, shutil
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK

RESULTS = '/home/lirank/rv-npu-jimu/rv-npu-jimu.v10/jimu-dse/results'
SRC_DOCX = RESULTS + '/NPU固件优化重构报告.docx'
SRC_MD   = RESULTS + '/平台化改造规划.md'
TARGET   = RESULTS + '/NPU固件优化与平台化改造报告.docx'

# start from the existing rendered report so ALL of its content is preserved
shutil.copyfile(SRC_DOCX, TARGET)
doc = Document(TARGET)
body = doc.element.body

# ---- remember trailing sectPr (keep it in body so add_table can read sections) ----
sectPr = None
for child in list(body):
    if child.tag == qn('w:sectPr'):
        sectPr = child

# ---------- formatting helpers (identical to build_report.py) ----------
def _run(r, latin='Arial', east='宋体', size=11, bold=False, color=None):
    r.font.size = Pt(size); r.font.bold = bold
    rpr = r._r.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    rf.set(qn('w:ascii'), latin); rf.set(qn('w:hAnsi'), latin); rf.set(qn('w:eastAsia'), east)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)

def _shade(p, fill='F2F2F2'):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def TITLE(text):
    try:
        p = doc.add_paragraph(style='Heading 1')
    except Exception:
        p = doc.add_paragraph()
    r = p.add_run(text); _run(r, size=18, bold=True, color='1F3864'); return p
def H2(text, page_break=False, size=15):
    p = doc.add_paragraph()
    if page_break:
        p.add_run().add_break(WD_BREAK.PAGE)
    r = p.add_run(text); _run(r, size=size, bold=True, color='1F3864'); return p
def H3(text, size=13):
    p = doc.add_paragraph(); r = p.add_run(text); _run(r, size=size, bold=True, color='2E5395'); return p
def H4(text, size=11.5):
    p = doc.add_paragraph(); r = p.add_run(text); _run(r, size=size, bold=True); return p
def BODY(text, indent=0):
    p = doc.add_paragraph(); r = p.add_run(text); _run(r, size=11)
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    return p
def CODE(text, fill='F2F2F2'):
    p = doc.add_paragraph(); _shade(p, fill)
    lines = text.split('\n')
    for i, ln in enumerate(lines):
        r = p.add_run(ln); _run(r, latin='Consolas', east='宋体', size=9)
        if i < len(lines) - 1:
            r._r.append(OxmlElement('w:br'))
    return p
def TABLE(rows):
    ncol = max(len(r) for r in rows)
    rows = [r + [''] * (ncol - len(r)) for r in rows]
    t = doc.add_table(rows=len(rows), cols=ncol)
    try: t.style = doc.styles['_Style 10']
    except Exception:
        try: t.style = doc.styles['Table Grid']
        except Exception: pass
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci); cell.text = ''
            r = cell.paragraphs[0].add_run(val); _run(r, size=10, bold=(ri == 0))
    return t

# ---------- inline markdown cleanup ----------
def _clean(s):
    s = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', s)   # [label](url) -> label
    s = s.replace('**', '').replace('`', '')          # drop bold / code markers
    return s.strip()

def _is_table_sep(cells):
    return all(re.fullmatch(r':?-{2,}:?', c.strip() or '') for c in cells) and len(cells) > 0

def _split_row(line):
    line = line.strip()
    if line.startswith('|'): line = line[1:]
    if line.endswith('|'):   line = line[:-1]
    return [c.strip() for c in line.split('|')]

# ---------- markdown renderer ----------
def render_markdown(md):
    lines = md.split('\n')
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # code fence
        if stripped.startswith('```'):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith('```'):
                buf.append(lines[i]); i += 1
            i += 1  # skip closing fence
            CODE('\n'.join(buf) if buf else ' ')
            continue

        # table (line starts with | and next line is a separator)
        if stripped.startswith('|') and i + 1 < n and _is_table_sep(_split_row(lines[i + 1])):
            header = [_clean(c) for c in _split_row(line)]
            i += 2  # skip header + separator
            rows = [header]
            while i < n and lines[i].strip().startswith('|'):
                rows.append([_clean(c) for c in _split_row(lines[i])])
                i += 1
            TABLE(rows)
            continue

        # headings
        m = re.match(r'(#{1,6})\s+(.*)', stripped)
        if m:
            level = len(m.group(1)); text = _clean(m.group(2))
            if level == 1:      H3(text, size=14)
            elif level == 2:    H3(text)
            else:               H4(text)
            i += 1
            continue

        # horizontal rule -> skip
        if re.fullmatch(r'(-{3,}|\*{3,}|_{3,})', stripped):
            i += 1
            continue

        # blockquote
        if stripped.startswith('>'):
            buf = []
            while i < n and lines[i].strip().startswith('>'):
                buf.append(re.sub(r'^\s*>\s?', '', lines[i])); i += 1
            BODY(_clean(' '.join(b for b in buf if b.strip())), indent=12)
            continue

        # list item (-, *, or "N.")
        lm = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)', line)
        if lm:
            indent_spaces = len(lm.group(1))
            marker = lm.group(2)
            bullet = '• ' if marker in ('-', '*') else marker + ' '
            BODY(bullet + _clean(lm.group(3)), indent=12 + indent_spaces * 3)
            i += 1
            continue

        # blank line
        if not stripped:
            i += 1
            continue

        # normal paragraph
        BODY(_clean(stripped))
        i += 1

# ============================================================
# bridge section + rendered planning doc
# ============================================================
H2('五、平台化改造规划（脱离 BERT 负载 + 算子级优化 + 硬件协同）', page_break=True)
BODY('本部分为在上述两轮优化结果基础上形成的平台化改造规划，完整保留自《平台化改造规划》'
     '文档，讲述如何让优化系统脱离 BERT 负载绑定、支持多算子固件优化，并反向驱动 NPU 硬件设计。')

with open(SRC_MD, encoding='utf-8') as f:
    md_text = f.read()
render_markdown(md_text)

# ---- ensure sectPr stays last (content was appended after it) ----
if sectPr is not None:
    body.remove(sectPr)
    body.append(sectPr)

doc.save(TARGET)
print('SAVED:', TARGET)
print('paragraphs:', len(doc.paragraphs), 'tables:', len(doc.tables))
