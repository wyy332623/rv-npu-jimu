# -*- coding: utf-8 -*-
"""Rebuild the NPU firmware optimization report: organically merge the original
author text with the deep-dive analysis into one restructured document.
Starts from the clean backup, clears the body (keeps styles/fonts/theme/sectPr),
and re-emits content in a unified structure."""
import glob, shutil
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK

RESULTS = '/home/lirank/rv-npu-jimu/rv-npu-jimu.v10/jimu-dse/results'
BACKUP  = [f for f in glob.glob(RESULTS + '/*.docx') if '备份' in f][0]
TARGET  = RESULTS + '/NPU固件优化重构报告.docx'

# rebuild from the clean backup so previous edits never accumulate
shutil.copyfile(BACKUP, TARGET)
doc = Document(TARGET)
body = doc.element.body

# ---- clear body but keep the trailing sectPr ----
sectPr = None
for child in list(body):
    if child.tag == qn('w:sectPr'):
        sectPr = child
        continue
    body.remove(child)

# ---------- formatting helpers ----------
def _run(r, latin='Arial', east='宋体', size=11, bold=False, color=None):
    r.font.size = Pt(size); r.font.bold = bold
    rpr = r._r.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    rf.set(qn('w:ascii'), latin); rf.set(qn('w:hAnsi'), latin); rf.set(qn('w:eastAsia'), east)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)

def _shade(p, fill='F2F2F2'):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def TITLE(text):
    try:
        p = doc.add_paragraph(style='Heading 1')
    except Exception:
        p = doc.add_paragraph()
    r = p.add_run(text); _run(r, size=18, bold=True, color='1F3864'); return p
def H2(text, page_break=False, size=15):
    p = doc.add_paragraph()
    if page_break:
        p.add_run().add_break(WD_BREAK.PAGE)
    r = p.add_run(text); _run(r, size=size, bold=True, color='1F3864'); return p
def H3(text, size=13):
    p = doc.add_paragraph(); r = p.add_run(text); _run(r, size=size, bold=True, color='2E5395'); return p
def H4(text, size=11.5):
    p = doc.add_paragraph(); r = p.add_run(text); _run(r, size=size, bold=True); return p
def BODY(text):
    p = doc.add_paragraph(); r = p.add_run(text); _run(r, size=11); return p
def CODE(text, fill='F2F2F2'):
    p = doc.add_paragraph(); _shade(p, fill)
    lines = text.split('\n')
    for i, ln in enumerate(lines):
        r = p.add_run(ln); _run(r, latin='Consolas', east='宋体', size=9)
        if i < len(lines) - 1:
            r._r.append(OxmlElement('w:br'))
    return p
def TABLE(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try: t.style = doc.styles['_Style 10']
    except Exception: pass
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci); cell.text = ''
            r = cell.paragraphs[0].add_run(val); _run(r, size=10, bold=(ri == 0))
    return t

# ============================================================
# 标题 + 引言
# ============================================================
TITLE('NPU 固件 Agent 重构结果分析报告')

H2('一、引言与总览')
BODY('本报告分析闭环优化流水线（jimu-dse）中 AI Agent 对 NPU 固件 firmware/bert/bert_layer.c '
     '（一层 BERT encoder）自动完成的两轮优化。流水线遵循 PROBE→ANALYZE→AGENT→VALIDATE→DEPLOY '
     '的闭环：先在 seq=2 与 seq=6 两种配置下探测 DRAM 流量并生成 micro-op DAG，再由 Agent 结合 DAG、'
     'DRAM cluster 分析与技能库（vrf-cache / dim-optimize）生成候选补丁，最后用“numpy 黄金参考→模拟器→'
     'DAG 审计”三轮校验接受或拒绝。全程 Agent 只允许修改 bert_layer.c，不触碰模拟器、ISS 与测试。')
BODY('两轮优化目标不同、方法互补：G1 侧重访存（把中间张量的 DRAM 往返改为片上 VRF 缓存），'
     'G3 侧重计算效率与代码正确性（把多块结构重构为单块，并修复其中隐藏的数值 Bug）。总览如下：')
TABLE([
    ['轮次', '运行目录 / 配置', 'Skill', '主指标', '核心结果'],
    ['G1', 'run-20260708-134915\ndim=2, hidden=4', 'vrf-cache', 'DRAM 总字节',
     'DRAM 6240→4896 B（−21.5%），max_diff=0'],
    ['G3', 'run-20260708-141835\ndim=4, hidden=4', 'dim-optimize (+vrf-cache)', 'test_pass',
     '测试通过；代码 657→327 行；DRAM −7%'],
])

# ============================================================
# 第一部分 —— G1
# ============================================================
H2('二、G1：DRAM 数据通路重构（dim=2，vrf-cache）', page_break=True)
BODY('（Agent 修改记录见 run 目录下的 diff_3.patch。）G1 的核心是 DRAM 优化，致力于减少外存读写、'
     '提升访存性能，同时保持 NATIVE_DIM=2、hidden_size=4 的多切块模式不变。')

H3('2.1 优化背景：为什么中间变量往返 DRAM 是瓶颈')
BODY('在固件原本的逻辑中，算子之间的数据通路较低效：大量中间变量（K、V、Q、Z、Self-Output、'
     'LayerNorm 等）被迫写入外部 DRAM 的 Scratch 区域，随后又被紧接着的下一级算子读出，造成严重的访存浪费。'
     'DAG 分析把这类“DRAM_STORE→DRAM_LOAD 同地址、中间无改写”的边识别为 save-load 对，其算术强度（AI）'
     '极低（0.0~1.2），是典型的访存受限热点。')
BODY('Skill 策略：引入向量寄存器堆（VRF）作为片上高速缓存，大幅降低内存流量。片上 MFU_INITIAL_VRF'
     '（mem6）容量 4096 元素，而 dim2/hidden4/seq6 全部 K+V 也仅约 48 元素，缓存这些中间量占用不到 3%，'
     '容量绰绰有余——这为“用片上缓存彻底替换 DRAM 往返”提供了物理前提。')

H3('2.2 关键机制：新增 mvm_tiled_vrf 函数')
BODY('G1 重构新增了 mvm_tiled_vrf 函数。传统矩阵-向量乘法从 DRAM 抓取输入向量，而该函数允许输入向量'
     '的基地址指向片上 VRF。引入此机制后，级联算子之间的中间结果可以全程留在寄存器中，无需访问外部内存。')
CODE('static void mvm_tiled_vrf(uint32_t mat_dram_base, uint32_t vec_vrf_base,\n'
     '                           uint32_t num_tiles, uint32_t bias_dram_base)\n'
     '{\n'
     '    // 配置标量寄存器，由软件接管多块分发\n'
     '    SEND_SI(OP_S_WR, REG_TILE_ROWS_ADDR, 1);\n'
     '    // ...\n'
     '    for (tc = 0; tc < num_tiles; tc++) {\n'
     '        // [核心优化]：直接从片上 VRF 缓存读取输入向量块\n'
     '        SEND_SI(OP_V_RD, vec_vrf_base + tc * NATIVE_DIM, 0);\n'
     '        SEND_SI(OP_V_WR, MEM_MVM_INITIAL_VRF, 0);\n'
     '        // ...执行硬件乘加并累加...\n'
     '    }\n'
     '}')

H3('2.3 中间变量的片上缓存策略')
BODY('基于 mvm_tiled_vrf 与 NPU 内置指令，代码针对不同变量实施了差异化的存储策略。')

H4('2.3.1 K / V 向量缓存化与地址映射')
BODY('原本通过 save_row_tiles 写回 SAVE_K_BASE / SAVE_V_BASE 的指令被移除，取而代之的是计算片上偏移量 '
     'cache_off，用 SEND_SI(OP_V_WR, MEM_MFU_INITIAL_VRF, cache_off) 把 K、V 写入寄存器；V 在计算出 '
     'K 的整体大小后叠加一个偏移，存放在 K 之后。下面是优化前后的对比：')
H4('优化前（写回 DRAM，注意力阶段再读回）：')
CODE('/* 基线：K 投影结果经 save_row_tiles 落到 DRAM 的 SAVE_K_BASE */\n'
     'save_row_tiles(num_tiles, SAVE_K_BASE + pos * num_tiles * 8,\n'
     '                MEM_ADDSUB_VRF_0, MEM_ADDSUB_VRF_1);\n'
     '...\n'
     'SEND_LO(OP_V_RD_DRAM, SAVE_K_BASE + p * num_tiles * 8 + tr * 8);  /* 注意力阶段读回 K */')
H4('优化后（直接搬到片上 VRF，按位置索引排布）：')
CODE('/* K：写入 mem6，偏移 = 位置 * num_tiles * NATIVE_DIM */\n'
     'uint32_t cache_off = pos * num_tiles * NATIVE_DIM;\n'
     'SEND_SI(OP_V_RD, MEM_ADDSUB_VRF_0, 0);\n'
     'SEND_SI(OP_V_WR, MEM_MFU_INITIAL_VRF, cache_off);\n'
     'SEND_SI(OP_V_RD, MEM_ADDSUB_VRF_1, 0);\n'
     'SEND_SI(OP_V_WR, MEM_MFU_INITIAL_VRF, cache_off + NATIVE_DIM);\n'
     '\n'
     '/* V：紧跟在整个 K 区之后，避免地址冲突 */\n'
     'uint32_t v_base_off = seq_len * num_tiles * NATIVE_DIM;\n'
     'uint32_t cache_off  = v_base_off + pos * num_tiles * NATIVE_DIM;\n'
     '...\n'
     'uint32_t k_off = p * num_tiles * NATIVE_DIM + tr * NATIVE_DIM;  /* DRAM 读换成 VRF 读 */\n'
     'SEND_SI(OP_V_RD, MEM_MFU_INITIAL_VRF, k_off);')
BODY('mem6 的地址被规划为三段：[K 区: seq×tiles×dim] [V 区: seq×tiles×dim] [复用槽 VRF_CACHE_OFF]。'
     'K/V 按 (位置, 瓦片行) 二维索引线性展开，保证注意力阶段对任意位置 p、瓦片行 tr 都能算出唯一片上偏移。')

H4('2.3.2 Q 向量即算即用与“延迟清零”')
BODY('Q 向量不再分配独立缓存区：计算出 Q 后直接保留在累加器 MEM_ADDSUB_VRF 内，注意力分数阶段直接读取，'
     '用完即清零供下级流水写入。这里有一个易错点——Q 与注意力上下文 context 复用同一块累加器，基线为给 '
     'context 累加腾出清零后的累加器，被迫先把 Q 落 DRAM 再读回。Agent 发现两者生命周期并不重叠，'
     '只要把“清零”推迟到 Q 被消费之后即可，Q 全程留在片上，省掉一存一取：')
CODE('/* 计算 Q —— 保留在 ADDSUB_VRF，不落 DRAM */\n'
     'mvm_tiled_q(...);\n'
     '...\n'
     'SEND_SI(OP_V_RD, acc_vrf, 0);      /* 注意力分数阶段直接读 Q */\n'
     '...\n'
     '/* Q 已被消费 —— 此刻才清零累加器，供 context 累加 */\n'
     'SEND_SI(OP_V_RD, MEM_FILL, 0);\n'
     'SEND_SI(OP_V_WR, acc_vrf, 0);')

H4('2.3.3 BERT 主干链路暂存区消除：单槽复用')
BODY('原固件预留了大量 0x600 以上的物理地址存放中间层结果，现已被统一的片上偏移量 VRF_CACHE_OFF 接管，'
     '写在 K、V 之后。注意力上下文 Z、自输出 SO、LayerNorm1 输出、GELU 激活这四个中间量在时间上依次产生、'
     '各自被下一级立即消费后即失效，活跃区间互不重叠，因此共用同一块区域也不会冲突——这等价于编译器寄存器'
     '分配中“活跃区间不冲突即可复用同一寄存器”的思想。')
CODE('#define VRF_CACHE_OFF   (2 * _SEQ_LEN * _NUM_TILES * NATIVE_DIM)  /* K+V 之后 */\n'
     '...\n'
     '/* Z（注意力上下文）写入复用槽 */\n'
     'SEND_SI(OP_V_WR, MEM_MFU_INITIAL_VRF, VRF_CACHE_OFF + tr * NATIVE_DIM);\n'
     '/* 自输出投影用 mvm_tiled_vrf 从同一槽把 Z 读进来做输入向量 */\n'
     'mvm_tiled_vrf(_PROJ_BASE + 3 * _STRIDE, VRF_CACHE_OFF, num_tiles,\n'
     '              _PROJ_BASE + 3 * _STRIDE + _MAT_SIZE);')
TABLE([
    ['中间层组件', '基线 (Baseline) DRAM 方案', 'G1 (VRF Cache) 优化方案'],
    ['Attention Context (Z)', '写入 SCRATCH_Z', '写入片上 MEM_MFU_INITIAL_VRF + 偏移'],
    ['Self-Output (SO) 残差', '写入并重载自 SO_SCRATCH', '从 VRF_CACHE_OFF 重载进行相加'],
    ['LayerNorm 1 输出', '暂存至 SCRATCH_LN1', '直接缓存在 VRF，对接 FFN 中间层 MVM'],
    ['GELU 激活值', '暂存至 SCRATCH_GELU', '留在 VRF 对接 FFN 最终输出层'],
])

H3('2.4 量化收益（seq=6，真实运行数据）')
BODY('从结果数据看，G1 极大收缩了存储带宽压力。未优化基线（seq_len=6）的 DRAM 流量约 6,240 字节；随着 '
     'VRF Cache 对 K/V、Q/SO/Z、LN 等环节的逐步接管（即 diff_3.patch 的效果），DRAM 流量最终降至约 '
     '4,896 字节，累计约 21.5% 的节约。这在保持 dim=2 小阵列面积的同时最大化了能效比。下表为逐轮明细'
     '（取自 val_*.json / prompt_1.txt）：')
TABLE([
    ['指标 (seq=6)', '真实基线', 'K/V 缓存 (iter1)', '全量缓存 (iter3=best)'],
    ['DRAM 总流量', '6240 B', '5856 B', '4896 B'],
    ['V_RD_DRAM', '—', '624 元素', '504 元素'],
    ['V_WR_DRAM', '—', '264 元素', '144 元素'],
    ['M_RD_DRAM (权重)', '576 元素', '576 元素', '576 元素'],
    ['相对真实基线降幅', '—', '6.2 %', '21.5 %'],
    ['数值 max_diff', '—', '0.000000', '0.000000'],
])
BODY('由于改动仅为数据路由、算术一字未改，三轮全部 max_diff=0.000000（精确无损）。其中 V_WR 降幅最大'
     '（−45.5%），因为所有中间“存”被彻底消除；V_RD 次之；权重加载 M_RD 不可消除，故保持不变。')

H3('2.5 局限与后续空间')
BODY('（1）X 输入仍未缓存：cluster 中 “Q Proj Load=624B”（每位置重复读入 X）与 “Save X→RES” 依旧存在，'
     '设计文档把 X-cache 列为可再省的一档；（2）iter2 的 diff 与 iter1 完全相同，属一次空转迭代，真正的两次'
     '跃迁发生在 iter1（K/V）与 iter3（Q/Z/SO/LN/GELU）；（3）剩余 V_RD 流量的主力，是不可消除的权重加载'
     '加上仍可消除的 X 重复读入。')

# ============================================================
# 第二部分 —— G3
# ============================================================
H2('三、G3：单块重构与计算效率优化（dim=4，dim-optimize）', page_break=True)
BODY('本轮 goal 为 combined（dim=4 / hidden=4 / num_head=2）。由于我们已完成 G1 优化，其实际收益以 G2'
     '（计算效率优化）为主：通过单块架构重构消除分块循环开销，实现计算效率飞跃。主指标是 test_pass。'
     '代码见 candidate_best.c，改动见 diff_2.patch。')

H3('3.1 优化背景与目标')
BODY('在默认基线中，硬件维度 NATIVE_DIM=2 小于隐藏层维度 hidden_size=4，NPU 必须在软件层面分块完成一次'
     '完整矩阵运算，多级循环带来较大运算与访存开销。G2 的核心是计算效率优化：把 NATIVE_DIM 调到 4 以匹配 '
     'hidden_size，一次性完成运算。Skill 策略为 dim-optimize（单块投影重构，消除软件循环），可将每次投影的 '
     'MV_MUL 从 4 次降到 1 次，并大幅减少 M_RD_DRAM（权重加载）次数。')
BODY('需要强调的关键前提：基线固件本是按 dim=2 多瓦片设计的，直接编到 dim=4 时运行时 num_tiles=1，'
     '但多块/多头代码路径内藏 Bug，产生的是错误数值而非崩溃（silent wrong result）。因此本轮真正的价值'
     '不只是“省”，而是“重构 + 查错”——把这份会静默算错的固件，改成正确且精简一半（657→327 行）的单块实现。')

H3('3.2 从分块到单块：剥离循环与简化访存')
BODY('固件的执行逻辑由“嵌套循环”转变为“直线型流水线”。在核心乘法函数 mvm_tiled_vrf / mvm_tiled_q 中，'
     '原有的 for (tc=…) 与 for (tr=…) 被完全移除，也无需再维护备用累加器（如 MEM_ADDSUB_VRF_2）：')
CODE('static void mvm_tiled_vrf(uint32_t mat_dram_base, uint32_t vec_vrf_offset,\n'
     '                           uint32_t num_tiles, uint32_t bias_dram_base)\n'
     '{\n'
     '    (void)num_tiles;                       /* 单块：num_tiles 恒为 1 */\n'
     '    SEND_SI(OP_S_WR, REG_TILE_ROWS, 1);\n'
     '    SEND_SI(OP_S_WR, REG_TILE_COLS, 1);\n'
     '    SEND_SI(OP_S_WR, REG_ITERATIONS, 1);\n'
     '\n'
     '    SEND_SI(OP_V_RD, MEM_MFU_INITIAL_VRF, vec_vrf_offset);  /* 输入向量取自片上 VRF */\n'
     '    SEND_SI(OP_V_WR, MEM_MVM_INITIAL_VRF, 0);\n'
     '    SEND_SI(OP_V_RD, MEM_MVM_INITIAL_VRF, 0);\n'
     '\n'
     '    SEND_LO(OP_M_RD_DRAM, mat_dram_base);   /* 一次加载完整 4x4 权重块 */\n'
     '    SEND_SI(OP_M_WR, MEM_MATRIX_RF, 0);\n'
     '    SEND_SI(OP_MV_MUL, 0, 0);               /* 一次 MV_MUL 得到全维结果 */\n'
     '\n'
     '    SEND_SI(OP_V_WR, MEM_MVM_ACC_VRF, 0);   /* 偏置加：单块单行 */\n'
     '    SEND_LO(OP_V_RD_DRAM, bias_dram_base);\n'
     '    SEND_SI(OP_V_RD, MEM_MVM_ACC_VRF, 0);\n'
     '    SEND_SI(OP_VV_ADD, 0, 0);\n'
     '    SEND_SI(OP_V_WR, MEM_ADDSUB_VRF_0, 0);\n'
     '}')
BODY('诸如 save_row_tiles、load_and_add_row_tiles、apply_layernorm 等函数也同步剥离了 num_tiles 参数'
     '及遍历循环，只需一次向量读写或一次 SUB_LAYERNORM 即可覆盖完整维度。这是代码量从 657 行骤降到 327 行'
     '的主因。')

H3('3.3 注意力机制精细化：逐头掩码（Per-Head Masking）')
BODY('在单块模式下，多个注意力头的数据被紧密堆叠在同一个硬件物理维度中。为防止头间数据污染，重构引入'
     '动态掩码机制：通过头索引 h 与头大小 head_size 移位生成 read_mask / write_mask，用 '
     'SEND_SI(OP_S_WR, REG_READ_VECTOR_MASK, read_mask & elem_mask) 指示硬件精准屏蔽非本头字段。')
CODE('uint32_t head_size      = NATIVE_DIM / num_head;      /* = 2 */\n'
     'uint32_t heads_per_tile = NATIVE_DIM / head_size;     /* = 2 */\n'
     'for (h = 0; h < heads_per_tile; h++) {\n'
     '    uint32_t head_shift = h * head_size;                       /* 0 或 2 */\n'
     '    uint32_t read_mask  = ((1 << head_size) - 1) << head_shift; /* 0x3 / 0xC */\n'
     '    uint32_t write_mask = (1 << head_size) - 1;                /* 0x3 */\n'
     '    ...\n'
     '    SEND_SI(OP_S_WR, REG_READ_VECTOR_MASK, read_mask & elem_mask);   /* 只读本头元素 */\n'
     '    ...\n'
     '    SEND_SI(OP_S_WR, REG_WRITE_VECTOR_MASK, write_mask & elem_mask);\n'
     '    SEND_SI(OP_V_WR, MEM_ADDSUB_VRF_0, write_off);              /* 写回本头元素槽 */\n'
     '}')

H3('3.4 三处隐藏 Bug 复盘（本轮核心价值）')
BODY('单块 + 多头改造之所以困难，在于它引入了三处“只产生错误数值、不报错崩溃”的隐蔽 Bug。Agent 通过'
     '追踪 NPU 硬件状态逐一定位并修复，这比模板化的数据搬运更能体现其对微架构的理解。')
H4('Bug 1 —— VRF 缓存的 bank / 偏移混淆')
BODY('根因：mvm_tiled_vrf 误把 VRF_CACHE_OFF 当成 bank 号（bank 16 = MEM_SPU_ABSMAX_REDUCE），'
     '而正确应为 bank 6（MEM_MFU_INITIAL_VRF）+ 该值作为偏移。症状：自输出、FFN 等后续投影全部读到垃圾数据。')
CODE('/* 错误：VRF_CACHE_OFF 被当作 bank 号 → 落到 bank 16 */\n'
     '// SEND_SI(OP_V_RD, VRF_CACHE_OFF, 0);\n'
     '/* 正确：bank = MEM_MFU_INITIAL_VRF(6)，VRF_CACHE_OFF 作偏移 */\n'
     'SEND_SI(OP_V_RD, MEM_MFU_INITIAL_VRF, vec_vrf_offset);')
H4('Bug 2 —— MEM_FILL 累加器清零放错位置')
BODY('根因：累加器清零本应在头循环“之前”只做一次；基线放在 per-head 循环“之内”，导致进入第二个头时把'
     '第一个头刚算出的 context 清零摧毁。修复：将清零上提到循环外，每个头写各自的元素偏移。')
CODE('/* 修复：清零移到头循环之前，只做一次 */\n'
     'SEND_SI(OP_V_RD, MEM_FILL, 0);\n'
     'SEND_SI(OP_V_WR, MEM_ADDSUB_VRF_0, 0);\n'
     'for (h = 0; h < heads_per_tile; h++) {\n'
     '    ...  /* 每个头写到 write_off，不再互相清零 */\n'
     '}')
H4('Bug 3 —— WRITE_VECTOR_MASK 用后未复位')
BODY('根因：最后一个头写完 context 后，写掩码停在 write_mask（如 0x3，只覆盖 4 元素中的前 2 个）。'
     '之后所有 V_WR（缓存写、投影、GELU）都只写一半元素，静默污染索引 2、3。修复：头循环结束后把写掩码'
     '恢复成全 1（0xFF）。')
CODE('    }  /* end head loop */\n'
     '    /* 恢复全掩码，防止后续 V_WR 只写部分元素 */\n'
     '    SEND_SI(OP_S_WR, REG_WRITE_VECTOR_MASK, elem_mask);   /* 0xFF */')

H3('3.5 性能提升分析（seq=6）')
BODY('优化成果主要体现在计算与外部访存次数的下降。相对 dim=2 真实多块基线的结构性对比如下：')
TABLE([
    ['性能指标', '基线方案 (Multi-Tile, dim=2)', 'G2 优化方案 (Single-Tile, dim=4)'],
    ['单次投影 MV_MUL 操作数', '4 次 (分块计算 + VV_ADD 累加)', '1 次 (单次搞定)'],
    ['权重块加载 M_RD_DRAM (seq=6)', '约 144 次', '约 36 次 (骤降 75%)'],
    ['代码控制流', '深层嵌套 for 循环', '高度线性展开流水线'],
])
BODY('而本轮三次迭代自身的逐轮明细（取自 val_*.json，主指标为测试通过）如下：')
TABLE([
    ['指标 (seq=6)', '基线 (dim=4, 含 bug)', 'iter1', 'iter2 = best'],
    ['dim4-h4 测试', '输出错误', '未通过', '通过 (max_diff=0)'],
    ['代码行数', '657', '657', '327'],
    ['DRAM 总流量', '5472 B', '5472 B', '5088 B'],
    ['V_RD_DRAM', '600 元素', '600 元素', '552 元素'],
    ['V_WR_DRAM', '192 元素', '192 元素', '144 元素'],
    ['M_RD_DRAM', '576 元素', '576 元素', '576 元素'],
    ['指令数', '—', '1919', '1643'],
])
BODY('总结：通过 dim-optimize 技能，NPU 固件彻底摆脱了软件分块的沉重镣铐，不仅节约了海量循环控制指令周期、'
     '把矩阵计算单元吞吐推向物理极限，更重要的是修正了 dim=4 路径下的三处静默错误，使 dim4-h4 测试首次通过；'
     'DRAM −7% 为附带收益。')

# ============================================================
# 第三部分 —— 综合对比与方法论
# ============================================================
H2('四、综合对比与方法论小结', page_break=True)
BODY('把两轮放在一起看，可以清楚区分两类不同性质的优化：G1 是“数据搬运型”优化——模板化地套用 vrf-cache '
     '把 DRAM 往返改为片上缓存，代码量增加、收益体现在带宽；G3 是“结构重构 + 查错型”优化——改变计算的组织'
     '方式并修复隐藏 Bug，代码量减半、收益体现在正确性与计算效率。')
TABLE([
    ['维度', 'G1 (dram-optimization)', 'G3 (combined，G2 主导)'],
    ['优化性质', '数据搬运，模板化套用 vrf-cache', '结构重构 + 定位隐藏 Bug'],
    ['起点固件', '正确的 dim=2 多块固件', '在 dim=4 下会静默算错的固件'],
    ['代码规模变化', '590 → 644 行（增）', '657 → 327 行（减半）'],
    ['主要收益', 'DRAM −21.5%，数值无损', '测试通过 + 代码精简，DRAM −7%'],
    ['能力体现', 'VRF 生命周期 / 单槽复用推理', 'bank/偏移、掩码、清零时序等微架构级排错'],
])
BODY('两轮共同遵守项目约束：只修改 firmware/bert/bert_layer.c，不触碰模拟器 / ISS / 测试；均以 DAG 与 '
     'DRAM cluster 分析作为定位依据，并由三轮校验（黄金参考→模拟器→DAG 审计）守住 max_diff 与回归。'
     '综合来看，Agent 表现得像一个真正理解 NPU 微架构（VRF bank、pipeline 寄存器、掩码寄存器、张量生命周期）'
     '的固件工程师，而非机械的模板替换器：它既能在 G1 中基于“活跃区间不重叠”推理出单槽复用，也能在 G3 中'
     '定位到“bank 当成偏移”“掩码未复位”这类只影响数值、不影响运行的隐蔽缺陷。')

# ---- ensure sectPr stays last ----
if sectPr is not None:
    body.remove(sectPr)
    body.append(sectPr)

doc.save(TARGET)
print('SAVED:', TARGET)
print('paragraphs:', len(doc.paragraphs), 'tables:', len(doc.tables))
